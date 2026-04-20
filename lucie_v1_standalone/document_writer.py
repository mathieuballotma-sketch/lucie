"""
DocumentWriter — convertit une réponse markdown produite par le Rédacteur
en un fichier .docx propre utilisable par un avocat.

Conversion minimale :
  # titre        → Heading 1
  ## sous-titre  → Heading 2
  ### …          → Heading 3
  - item         → puce
  **gras**       → gras inline
  *italique*     → italique inline
  ligne vide     → séparation paragraphe
  autre          → paragraphe normal

Nom de fichier : {kind}_{slug}_{YYYY-MM-DD}.docx

Aucune dépendance au reste du repo ; dépend uniquement de `python-docx`.
"""

from __future__ import annotations

import re
import unicodedata
from datetime import date
from pathlib import Path
from typing import Iterable, Optional

from docx import Document

# Dossier de sortie par défaut (relatif au cwd d'exécution). Aligné sur
# `app.actions.word_output_dir` / `app.docs_dir` définis dans config.
_DEFAULT_OUTPUT_DIR = Path("./Lucid_Docs")

# Inline formatting : gras **…** puis italique *…* (ordre important).
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_INLINE_CODE_RE = re.compile(r"`([^`]+)`")

# Nettoyage du slug : on garde lettres, chiffres, tirets.
_SLUG_RE = re.compile(r"[^a-z0-9]+")


def _slugify(text: str, max_len: int = 40) -> str:
    """Produit un slug ASCII minuscule pour nom de fichier."""
    normalized = unicodedata.normalize("NFKD", text)
    ascii_text = normalized.encode("ascii", "ignore").decode("ascii").lower()
    slug = _SLUG_RE.sub("-", ascii_text).strip("-")
    if len(slug) > max_len:
        slug = slug[:max_len].rstrip("-")
    return slug or "document"


def _extract_title(markdown_text: str, fallback: str) -> str:
    """Extrait le premier titre `# …` du markdown pour le slug. À défaut `fallback`."""
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# ") and not stripped.startswith("## "):
            return stripped[2:].strip()
    return fallback


def _add_inline_runs(paragraph, text: str) -> None:
    """Ajoute des runs au paragraphe en interprétant **gras**, *italique*, `code`.

    Parsing simple basé sur tokenisation : on découpe sur chaque pattern successivement.
    Si ça se chevauche (rare), le gras prime.
    """
    # Tokens = liste de (text, bold, italic, code)
    tokens = [(text, False, False, False)]

    def _expand(pattern: re.Pattern[str], flag_index: int) -> None:
        nonlocal tokens
        new: list[tuple[str, bool, bool, bool]] = []
        for seg_text, b, i, c in tokens:
            if b or c:  # pas de re-parsing à l'intérieur d'un segment déjà stylé
                new.append((seg_text, b, i, c))
                continue
            last = 0
            for m in pattern.finditer(seg_text):
                if m.start() > last:
                    new.append((seg_text[last:m.start()], b, i, c))
                flags = [b, i, c]
                flags[flag_index] = True
                new.append((m.group(1), flags[0], flags[1], flags[2]))
                last = m.end()
            if last < len(seg_text):
                new.append((seg_text[last:], b, i, c))
        tokens = new

    _expand(_BOLD_RE, 0)
    _expand(_INLINE_CODE_RE, 2)
    _expand(_ITALIC_RE, 1)

    for seg_text, bold, italic, code in tokens:
        if not seg_text:
            continue
        run = paragraph.add_run(seg_text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if code:
            run.font.name = "Menlo"


def _iter_logical_lines(markdown_text: str) -> Iterable[str]:
    """Itère sur les lignes du markdown en ignorant un bloc final de disclaimer auto
    (ajouté par `pipeline._format_final`) — l'utilisateur n'a pas besoin du bloc
    « Note générée par Lucie V1 — Score de fiabilité… » dans son DOCX.
    """
    lines = markdown_text.splitlines()
    # Retirer disclaimer auto (tout après une ligne "---" en fin)
    separators = [i for i, line in enumerate(lines) if line.strip() == "---"]
    if separators and separators[-1] >= len(lines) - 10:
        lines = lines[: separators[-1]]
    for line in lines:
        yield line


def write_docx(
    markdown_text: str,
    kind: str = "document",
    output_dir: Optional[Path] = None,
    slug_hint: Optional[str] = None,
) -> Path:
    """Convertit `markdown_text` en fichier .docx dans `output_dir`.

    Args:
        markdown_text: Texte markdown produit par le Rédacteur.
        kind: "courrier", "acte", "synthese", "note", ou "document".
        output_dir: Dossier cible. Défaut : ./Lucid_Docs (créé si nécessaire).
        slug_hint: Optionnel — texte source pour le slug si aucun titre `#`.

    Returns:
        Chemin absolu du fichier produit.
    """
    target_dir = Path(output_dir) if output_dir else _DEFAULT_OUTPUT_DIR
    target_dir.mkdir(parents=True, exist_ok=True)

    title = _extract_title(markdown_text, fallback=slug_hint or kind)
    slug = _slugify(title)
    filename = f"{kind}_{slug}_{date.today().isoformat()}.docx"
    path = (target_dir / filename).resolve()

    doc = Document()

    for raw in _iter_logical_lines(markdown_text):
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            # Saut de paragraphe : on ne crée rien, python-docx gère naturellement.
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, stripped[2:].strip())
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            # strip "1. " prefix
            body = re.sub(r"^\d+\.\s+", "", stripped)
            _add_inline_runs(p, body)
        elif stripped.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            _add_inline_runs(p, stripped[2:].strip())
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, stripped)

    doc.save(str(path))
    return path
