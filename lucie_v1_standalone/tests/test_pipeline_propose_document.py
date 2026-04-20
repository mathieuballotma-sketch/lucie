"""
Tests pour les améliorations UX du pipeline :
  - Proposition avant production (EXPLICIT_ORDER + verbe → produces_document)
  - Boutons Oui/Non génériques sur questions fermées (suggested_replies)
  - Routing de la décision utilisateur (__decision__:…)
  - Helpers de détection (_detect_production_request, _extract_decision)
  - document_writer.write_docx (roundtrip markdown → DOCX)

Ces tests sont 0 LLM et tournent sans Ollama.

Lancer avec :
    python -m pytest lucie_v1_standalone/tests/test_pipeline_propose_document.py -v
"""
from __future__ import annotations

import tempfile
from pathlib import Path

import pytest

from lucie_v1_standalone import document_writer
from lucie_v1_standalone import pipeline
from lucie_v1_standalone.pipeline import (
    PipelineResponse,
    _attach_suggested_replies,
    _build_proposition,
    _detect_production_request,
    _extract_decision,
)


# ─── Helpers pipeline (0 LLM, purement synchrones) ──────────────────────────

class TestExtractDecision:
    def test_no_marker_returns_none_and_untouched_query(self):
        decision, original = _extract_decision("Quelle est la loi applicable ?")
        assert decision is None
        assert original == "Quelle est la loi applicable ?"

    def test_yes_produce_marker_parses_value_and_original(self):
        decision, original = _extract_decision(
            "__decision__:yes_produce|original=Rédige une mise en demeure"
        )
        assert decision == "yes_produce"
        assert original == "Rédige une mise en demeure"

    def test_no_marker_parses(self):
        decision, original = _extract_decision(
            "__decision__:no|original=Voulez-vous continuer ?"
        )
        assert decision == "no"
        assert original == "Voulez-vous continuer ?"

    def test_marker_without_original_returns_empty_original(self):
        decision, original = _extract_decision("__decision__:yes_produce")
        assert decision == "yes_produce"
        assert original == ""


class TestDetectProductionRequest:
    @pytest.mark.parametrize("query,expected_kind", [
        ("Rédige un projet de courrier de mise en demeure", "courrier"),
        ("Prépare une lettre de licenciement", "courrier"),
        ("Écris-moi une synthèse juridique", "synthese"),
        ("Projet de conclusions pour le 15 mai", "acte"),
        ("Prépare une note d'analyse", "note"),
        ("Génère un email de réponse", "courrier"),
    ])
    def test_production_verbs_and_kinds(self, query, expected_kind):
        assert _detect_production_request(query) == expected_kind

    @pytest.mark.parametrize("query", [
        "Quel est l'article 1231-1 du Code civil ?",
        "Comment calcule-t-on l'indemnité de licenciement ?",
        "Bonjour Lucie",
        "Explique-moi la procédure de licenciement économique",
    ])
    def test_factual_queries_return_none(self, query):
        assert _detect_production_request(query) is None

    def test_production_verb_without_kind_falls_back_to_document(self):
        # Verbe détecté mais aucun mot-clé de kind reconnu.
        assert _detect_production_request("Rédige quelque chose pour moi") == "document"


class TestBuildProposition:
    def test_courrier_proposition_shape(self):
        resp = _build_proposition("Rédige une mise en demeure", "courrier")
        assert isinstance(resp, PipelineResponse)
        assert resp.produces_document is True
        assert resp.document_kind == "courrier"
        assert resp.mode == "proposition"
        assert "courrier" in resp.answer
        assert "?" in resp.answer
        assert len(resp.suggested_replies) == 2
        assert resp.suggested_replies[0]["value"] == "yes_produce"
        assert resp.suggested_replies[1]["value"] == "no_text"

    def test_all_kinds_have_label(self):
        for kind in ("courrier", "acte", "synthese", "note", "document"):
            resp = _build_proposition("Prépare", kind)
            assert resp.answer  # chaque kind produit un texte non vide


class TestAttachSuggestedReplies:
    def test_closed_question_triggers_yes_no(self):
        r = PipelineResponse(answer="Voulez-vous que je précise le délai ?")
        _attach_suggested_replies(r)
        assert r.suggested_replies == [
            {"label": "Oui", "value": "yes"},
            {"label": "Non", "value": "no"},
        ]

    @pytest.mark.parametrize("text", [
        "Souhaitez-vous une analyse plus détaillée ?",
        "Dois-je ajouter une clause de confidentialité ?",
        "Confirmez-vous les éléments fournis ?",
        "Puis-je contacter l'adversaire pour vous ?",
        "Faut-il inclure les pièces justificatives ?",
    ])
    def test_various_closed_patterns(self, text):
        r = PipelineResponse(answer=text)
        _attach_suggested_replies(r)
        assert len(r.suggested_replies) == 2

    def test_no_closed_question_leaves_replies_empty(self):
        r = PipelineResponse(
            answer="L'article 1231-1 du Code civil dispose que le débiteur est tenu."
        )
        _attach_suggested_replies(r)
        assert r.suggested_replies == []

    def test_existing_replies_are_preserved(self):
        # Une proposition déjà construite ne doit pas être écrasée.
        r = PipelineResponse(
            answer="Voulez-vous que je produise ce courrier ?",
            suggested_replies=[
                {"label": "Oui, produire", "value": "yes_produce"},
                {"label": "Non", "value": "no_text"},
            ],
        )
        _attach_suggested_replies(r)
        assert r.suggested_replies[0]["value"] == "yes_produce"


# ─── Pipeline end-to-end (async, 0 LLM car SMALL_TALK/EXPLICIT_ORDER early) ──

@pytest.mark.asyncio
async def test_explicit_order_production_verb_returns_proposition():
    """Une demande de production ne lance PAS le pipeline complet — elle propose."""
    response = await pipeline.run("Rédige un projet de courrier de mise en demeure")
    assert response.produces_document is True
    assert response.document_kind == "courrier"
    assert response.mode == "proposition"
    assert len(response.suggested_replies) == 2
    # L'answer doit contenir une question à l'utilisateur.
    assert "?" in response.answer


@pytest.mark.asyncio
async def test_decision_no_text_returns_short_answer_without_pipeline():
    """Décision 'non' → réponse courte, pas de pipeline complet."""
    response = await pipeline.run(
        "__decision__:no_text|original=Rédige une mise en demeure"
    )
    assert response.produces_document is False
    assert response.document_path is None
    assert response.answer  # non vide
    # Pas de déclenchement LLM — la réponse "non" est déterministe.
    assert "disposition" in response.answer.lower() or "entendu" in response.answer.lower()


@pytest.mark.asyncio
async def test_small_talk_still_works_after_changes():
    """Non-régression : les salutations restent intactes."""
    response = await pipeline.run("Bonjour Lucie")
    assert response.verifier_score == 1.0
    assert response.produces_document is False
    assert response.suggested_replies == []


# ─── DocumentWriter ──────────────────────────────────────────────────────────

MARKDOWN_SAMPLE = """# Mise en demeure — Retard de paiement

## Contexte

Monsieur,

Conformément à l'**article 1231-1** du Code civil, je vous *mets en demeure* de régler la somme due.

## Montant réclamé

- Principal : 1 200 €
- Intérêts : 45 €
- Frais : 30 €

## Délai

1. Règlement sous 8 jours
2. À défaut, saisine du tribunal

> Cette lettre vaut mise en demeure formelle.

---
_Note générée par Lucie V1 — Score de fiabilité : 92% — Verdict : FIABLE_
_À vérifier par un avocat qualifié avant tout usage professionnel._
"""


class TestDocumentWriter:
    def test_write_docx_produces_file_with_expected_name(self, tmp_path):
        path = document_writer.write_docx(
            MARKDOWN_SAMPLE, kind="courrier", output_dir=tmp_path
        )
        assert path.exists()
        assert path.suffix == ".docx"
        assert path.name.startswith("courrier_")
        # Le slug doit venir du premier titre H1.
        assert "mise-en-demeure" in path.name
        # Le fichier DOCX a une taille minimale (header ZIP + xml).
        assert path.stat().st_size > 5_000

    def test_docx_contains_all_sections(self, tmp_path):
        from docx import Document
        path = document_writer.write_docx(
            MARKDOWN_SAMPLE, kind="courrier", output_dir=tmp_path
        )
        doc = Document(str(path))
        texts = [p.text for p in doc.paragraphs]
        # Les titres H1/H2 sont bien présents.
        assert any("Mise en demeure" in t for t in texts)
        assert any("Contexte" in t for t in texts)
        assert any("Montant" in t for t in texts)
        # Les puces sont présentes.
        assert any("Principal" in t for t in texts)
        # Le disclaimer auto en fin (après ---) doit AVOIR été coupé.
        assert not any("Score de fiabilité" in t for t in texts)

    def test_slug_fallback_when_no_h1(self, tmp_path):
        md = "Juste un paragraphe sans titre en tête."
        path = document_writer.write_docx(
            md, kind="note", output_dir=tmp_path, slug_hint="demande-ad-hoc"
        )
        assert path.exists()
        assert "demande-ad-hoc" in path.name

    def test_inline_bold_and_italic_runs_are_set(self, tmp_path):
        from docx import Document
        md = "Texte **en gras** et *en italique* et `code`."
        path = document_writer.write_docx(md, kind="note", output_dir=tmp_path)
        doc = Document(str(path))
        # Il existe au moins un run bold et un run italic dans les paragraphes.
        runs = [r for p in doc.paragraphs for r in p.runs]
        assert any(r.bold for r in runs)
        assert any(r.italic for r in runs)
