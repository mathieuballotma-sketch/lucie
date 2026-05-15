"""Tests unitaires du parser docx déterministe."""

from __future__ import annotations

import pytest
from docx import Document

from lucie_v1_standalone.document_analyzer.docx_parser import parse_docx
from lucie_v1_standalone.document_analyzer.exceptions import (
    CorruptedFileError,
    EmptyDocumentError,
)


def test_parse_docx_extracts_text(fixture_lic_perso_docx):
    text, pages = parse_docx(fixture_lic_perso_docx)
    assert pages >= 1
    assert "Paul Martin" in text
    assert "faute grave" in text
    assert "prud'hommes" in text.lower()


def test_parse_docx_paragraph_count_drives_pseudo_pages(tmp_path):
    """Avec 80 paragraphes, pseudo_pages = 80 // 40 = 2."""
    path = tmp_path / "long.docx"
    doc = Document()
    for i in range(80):
        doc.add_paragraph(f"Paragraphe {i} — texte de test.")
    doc.save(str(path))
    _text, pages = parse_docx(path)
    assert pages == 2


def test_parse_docx_empty_raises(tmp_path):
    path = tmp_path / "empty.docx"
    doc = Document()
    doc.save(str(path))
    with pytest.raises(EmptyDocumentError):
        parse_docx(path)


def test_parse_docx_corrupt_raises(tmp_path):
    path = tmp_path / "corrupt.docx"
    path.write_bytes(b"not a real docx")
    with pytest.raises(CorruptedFileError):
        parse_docx(path)


def test_parse_docx_missing_file_raises(tmp_path):
    with pytest.raises(CorruptedFileError):
        parse_docx(tmp_path / "does_not_exist.docx")
