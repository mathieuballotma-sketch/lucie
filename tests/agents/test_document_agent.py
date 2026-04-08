"""
Tests unitaires pour DocumentAgent — lecture PDF/DOCX + résumé LLM.

Pilier #2 BMAD : Lecture de Documents.

Couverture :
  - can_handle() avec requêtes de lecture (lis, résume, extrais, analyse…)
  - can_handle() avec requêtes de création (régression)
  - _tool_read_pdf() avec PDF multi-pages créé via fitz
  - _tool_read_docx() avec DOCX créé via python-docx
  - _tool_summarize_document() avec LLM mocké
  - Gestion des erreurs : fichier inexistant, format non supporté

Ces tests ne nécessitent PAS Ollama — le LLM est toujours mocké.
"""
from __future__ import annotations

import os
from pathlib import Path
from typing import Any
from unittest.mock import MagicMock

import pytest

from app.agents.document_agent import DocumentAgent


# ── Fixtures ───────────────────────────────────────────────────────────────────


@pytest.fixture
def agent() -> DocumentAgent:
    """DocumentAgent avec LLM mocké, sans Ollama."""
    llm = MagicMock()
    llm.generate = MagicMock(return_value="Résumé généré par le LLM mocké.")
    bus = MagicMock()
    config: dict[str, Any] = {}
    return DocumentAgent(llm, bus, config)


@pytest.fixture
def sample_pdf(tmp_path: Path) -> str:
    """PDF de test multi-pages créé via PyMuPDF (fitz) — pas de reportlab nécessaire."""
    import fitz  # PyMuPDF

    doc = fitz.open()
    page1 = doc.new_page()
    page1.insert_text((72, 72), "Page 1 : Contenu de test du PDF.")
    page1.insert_text((72, 100), "Ce document sert aux tests unitaires.")
    page2 = doc.new_page()
    page2.insert_text((72, 72), "Page 2 : Suite du document PDF de test.")
    pdf_path = str(tmp_path / "test.pdf")
    doc.save(pdf_path)
    doc.close()
    return pdf_path


@pytest.fixture
def sample_docx(tmp_path: Path) -> str:
    """DOCX de test créé via python-docx."""
    from docx import Document

    doc = Document()
    doc.add_heading("Titre du document de test", level=1)
    doc.add_paragraph("Premier paragraphe du document DOCX de test.")
    doc.add_paragraph("Deuxième paragraphe avec davantage de contenu.")
    docx_path = str(tmp_path / "test.docx")
    doc.save(docx_path)
    return docx_path


# ── can_handle : requêtes de LECTURE ──────────────────────────────────────────


def test_can_handle_lis_pdf(agent: DocumentAgent) -> None:
    assert agent.can_handle("lis ce pdf") is True


def test_can_handle_lire_document(agent: DocumentAgent) -> None:
    assert agent.can_handle("lire le document") is True


def test_can_handle_resume_le_document(agent: DocumentAgent) -> None:
    assert agent.can_handle("résume le document") is True


def test_can_handle_resumer_fichier(agent: DocumentAgent) -> None:
    assert agent.can_handle("résumer ce fichier pdf") is True


def test_can_handle_extrais_texte_pdf(agent: DocumentAgent) -> None:
    assert agent.can_handle("extrais le texte du pdf") is True


def test_can_handle_extraire_contenu(agent: DocumentAgent) -> None:
    assert agent.can_handle("extraire le contenu du document") is True


def test_can_handle_analyse_fichier_docx(agent: DocumentAgent) -> None:
    assert agent.can_handle("analyse ce fichier docx") is True


def test_can_handle_analyser_document(agent: DocumentAgent) -> None:
    assert agent.can_handle("analyser le document word") is True


def test_can_handle_ouvrir_word(agent: DocumentAgent) -> None:
    assert agent.can_handle("ouvrir le fichier word") is True


def test_can_handle_ouvre_pdf(agent: DocumentAgent) -> None:
    assert agent.can_handle("ouvre ce pdf s'il te plaît") is True


def test_can_handle_contenu_pdf(agent: DocumentAgent) -> None:
    assert agent.can_handle("montre le contenu du pdf") is True


def test_can_handle_texte_document(agent: DocumentAgent) -> None:
    assert agent.can_handle("quel est le texte de ce document") is True


# ── can_handle : requêtes de CRÉATION (régression) ─────────────────────────────


def test_can_handle_cree_word(agent: DocumentAgent) -> None:
    assert agent.can_handle("crée un document word") is True


def test_can_handle_genere_docx(agent: DocumentAgent) -> None:
    assert agent.can_handle("génère un fichier docx sur Python") is True


def test_can_handle_fais_document(agent: DocumentAgent) -> None:
    assert agent.can_handle("fais un document sur l'IA") is True


def test_can_handle_ecris_rapport(agent: DocumentAgent) -> None:
    assert agent.can_handle("écris un résumé en word") is True


def test_can_handle_redige_pdf(agent: DocumentAgent) -> None:
    assert agent.can_handle("rédige un document pdf") is True


# ── can_handle : faux positifs évités ─────────────────────────────────────────


def test_can_handle_no_match_question(agent: DocumentAgent) -> None:
    assert agent.can_handle("quelle heure est-il ?") is False


def test_can_handle_no_match_meteo(agent: DocumentAgent) -> None:
    assert agent.can_handle("quel temps fait-il à Paris ?") is False


def test_can_handle_no_doc_word(agent: DocumentAgent) -> None:
    # "lis" seul sans mot document → False
    assert agent.can_handle("lis-moi une blague") is False


# ── _tool_read_pdf ─────────────────────────────────────────────────────────────


@pytest.mark.anyio
async def test_read_pdf_contenu_page1(agent: DocumentAgent, sample_pdf: str) -> None:
    result = await agent._tool_read_pdf(path=sample_pdf)
    assert "Page 1" in result
    assert "Contenu de test" in result


@pytest.mark.anyio
async def test_read_pdf_multipage(agent: DocumentAgent, sample_pdf: str) -> None:
    result = await agent._tool_read_pdf(path=sample_pdf)
    assert "Page 2" in result
    assert "Suite du document" in result


@pytest.mark.anyio
async def test_read_pdf_header_present(agent: DocumentAgent, sample_pdf: str) -> None:
    result = await agent._tool_read_pdf(path=sample_pdf)
    assert "📄" in result
    assert "test.pdf" in result


@pytest.mark.anyio
async def test_read_pdf_fichier_introuvable(agent: DocumentAgent) -> None:
    result = await agent._tool_read_pdf(path="/nonexistent/path/fichier.pdf")
    assert "❌" in result
    assert "introuvable" in result.lower()


# ── _tool_read_docx ───────────────────────────────────────────────────────────


@pytest.mark.anyio
async def test_read_docx_contenu(agent: DocumentAgent, sample_docx: str) -> None:
    result = await agent._tool_read_docx(path=sample_docx)
    assert "Premier paragraphe" in result
    assert "Deuxième paragraphe" in result


@pytest.mark.anyio
async def test_read_docx_titre(agent: DocumentAgent, sample_docx: str) -> None:
    result = await agent._tool_read_docx(path=sample_docx)
    assert "Titre du document de test" in result


@pytest.mark.anyio
async def test_read_docx_header_present(agent: DocumentAgent, sample_docx: str) -> None:
    result = await agent._tool_read_docx(path=sample_docx)
    assert "📄" in result
    assert "test.docx" in result


@pytest.mark.anyio
async def test_read_docx_fichier_introuvable(agent: DocumentAgent) -> None:
    result = await agent._tool_read_docx(path="/nonexistent/path/fichier.docx")
    assert "❌" in result
    assert "introuvable" in result.lower()


# ── _tool_summarize_document ──────────────────────────────────────────────────


@pytest.mark.anyio
async def test_summarize_pdf_appelle_llm(agent: DocumentAgent, sample_pdf: str) -> None:
    agent.ask_llm = MagicMock(return_value="Résumé du document PDF.")
    result = await agent._tool_summarize_document(path=sample_pdf)
    agent.ask_llm.assert_called_once()
    assert "Résumé" in result


@pytest.mark.anyio
async def test_summarize_pdf_contient_nom_fichier(agent: DocumentAgent, sample_pdf: str) -> None:
    agent.ask_llm = MagicMock(return_value="Résumé concis.")
    result = await agent._tool_summarize_document(path=sample_pdf)
    assert "test.pdf" in result


@pytest.mark.anyio
async def test_summarize_docx_appelle_llm(agent: DocumentAgent, sample_docx: str) -> None:
    agent.ask_llm = MagicMock(return_value="Résumé du document DOCX.")
    result = await agent._tool_summarize_document(path=sample_docx)
    agent.ask_llm.assert_called_once()
    assert "Résumé" in result


@pytest.mark.anyio
async def test_summarize_docx_contient_nom_fichier(agent: DocumentAgent, sample_docx: str) -> None:
    agent.ask_llm = MagicMock(return_value="Résumé concis.")
    result = await agent._tool_summarize_document(path=sample_docx)
    assert "test.docx" in result


@pytest.mark.anyio
async def test_summarize_format_non_supporte(agent: DocumentAgent, tmp_path: Path) -> None:
    txt_path = str(tmp_path / "fichier.txt")
    Path(txt_path).write_text("contenu texte", encoding="utf-8")
    result = await agent._tool_summarize_document(path=txt_path)
    assert "❌" in result
    assert "non supporté" in result


@pytest.mark.anyio
async def test_summarize_fichier_introuvable(agent: DocumentAgent) -> None:
    result = await agent._tool_summarize_document(path="/nonexistent/fichier.pdf")
    assert "❌" in result
    assert "introuvable" in result.lower()


# ── get_tools : 4 outils enregistrés ─────────────────────────────────────────


def test_get_tools_count(agent: DocumentAgent) -> None:
    tools = agent.get_tools()
    assert len(tools) == 4


def test_get_tools_names(agent: DocumentAgent) -> None:
    names = {t.name for t in agent.get_tools()}
    assert names == {"create_word_document", "read_pdf", "read_docx", "summarize_document"}
