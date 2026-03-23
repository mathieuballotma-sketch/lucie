import os

from docx import Document

from ..utils.logger import get_logger
from ..utils.notifier import send_notification

logger = get_logger(__name__)


class WriterAgent:
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)
        logger.info(
            f"📄 WriterAgent initialisé, répertoire de sortie : " f"{self.output_dir}"
        )

    def create_word_document(self, title: str, content: str) -> str:
        try:
            # Valider le titre — rejeter les messages d'erreur utilisés comme titre
            error_markers = ("Erreur", "Error", "Exception", "Traceback")
            if any(marker in title for marker in error_markers):
                title = "document_sans_titre"
            # Tronquer à 80 caractères
            if len(title) > 80:
                title = title[:80]
            # Nettoyer le titre pour en faire un nom de fichier
            filename = title.replace(" ", "_")
            for ch in ('/', '\\', ':', '*', '"', '<', '>', '|', '?', '\n'):
                filename = filename.replace(ch, "_")
            # Supprimer les parenthèses et points de suspension
            filename = filename.replace("(", "").replace(")", "")
            if not filename or filename.strip("_") == "":
                filename = "document_sans_titre"
            filepath = os.path.join(self.output_dir, f"{filename}.docx")

            doc = Document()
            doc.add_heading(title, level=0)
            doc.add_paragraph(content)
            doc.save(filepath)

            logger.info(f"✅ Document Word créé : {filepath}")

            send_notification(
                title="Document créé", message=f"{filename}.docx", filepath=filepath
            )

            return (
                f"✅ Document Word créé : {filename}.docx "
                "(une notification a été envoyée)"
            )
        except Exception as e:
            logger.error(f"❌ Erreur lors de la création du document Word : {e}")
            return f"Erreur lors de la création du document Word : {str(e)}"
