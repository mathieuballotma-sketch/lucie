# app/actions/writer.py
import os
import subprocess

import AppKit
from docx import Document

from ..utils.logger import logger


class WriterAgent:
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)
        logger.info(f"📄 WriterAgent initialisé, répertoire de sortie : {
                self.output_dir}")

    def create_word_document(self, title: str, content: str) -> str:
        try:
            filename = title.replace(" ", "_").replace("/", "_").replace("\\", "_")
            if not filename:
                filename = "document"
            filepath = os.path.join(self.output_dir, f"{filename}.docx")

            doc = Document()
            doc.add_heading(title, level=0)
            doc.add_paragraph(content)
            doc.save(filepath)

            logger.info(f"✅ Document Word créé : {filepath}")

            # Notification interactive
            self._send_notification(
                title="Document créé", message=f"{filename}.docx", filepath=filepath
            )

            return f"✅ Document Word créé : {filename}.docx (une notification a été envoyée)"
        except Exception as e:
            logger.error(f"❌ Erreur lors de la création du document Word : {e}")
            return f"Erreur lors de la création du document Word : {str(e)}"

    def _send_notification(self, title: str, message: str, filepath: str):
        """Envoie une notification interactive avec NSUserNotification."""
        try:
            notification = AppKit.NSUserNotification.alloc().init()
            notification.setTitle_(title)
            notification.setInformativeText_(message)
            notification.setActionButtonTitle_("Ouvrir")
            notification.setUserInfo_({"filepath": filepath, "action": "open_file"})
            notification.setHasActionButton_(True)

            # Planifier la notification immédiatement
            centre = AppKit.NSUserNotificationCenter.defaultUserNotificationCenter()
            # Le délégué est l'AppDelegate
            centre.setDelegate_(AppKit.NSApp().delegate())
            centre.scheduleNotification_(notification)
        except Exception as e:
            logger.warning(f"Impossible d'envoyer la notification interactive : {e}")
            # Fallback sur notification simple
            try:
                script = f'display notification "{message}" with title "{title}"'
                subprocess.run(["osascript", "-e", script], capture_output=True)
            except BaseException:
                pass
